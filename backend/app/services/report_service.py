from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy.orm import Session

from backend.app.models.entities import ImageAsset, MetricRecord, RunTask
from backend.app.services.storage import storage_service


# ---------------------------------------------------------------------------
# 样式与字体配置
# ---------------------------------------------------------------------------

FONT_NAME_CN = "微软雅黑"
FONT_NAME_EN = "Times New Roman"


def _set_cell_bg(cell, hex_color: str) -> None:
    """设置单元格背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_table_row(table, cells_text: list[str], bold: bool = False, header: bool = False) -> None:
    """向表格追加一行。"""
    row = table.add_row()
    for idx, text in enumerate(cells_text):
        cell = row.cells[idx]
        cell.text = str(text) if text is not None else ""
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs
        if run:
            run[0].bold = bold or header
            run[0].font.name = FONT_NAME_EN
            run[0]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
            run[0].font.size = Pt(9)
        if header:
            _set_cell_bg(cell, "2E75B6")
            if run:
                run[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row


def _set_col_width(table, col_idx: int, width_cm: float) -> None:
    for cell in table.columns[col_idx].cells:
        cell.width = Cm(width_cm)


def _add_horizontal_rule(doc: Document) -> None:
    """添加一条水平分隔线（段落底边框）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_font(para, text: str, size_pt: int, bold: bool = False, color: RGBColor | None = None) -> None:
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _add_heading1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_paragraph_font(p, text, 16, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    _add_horizontal_rule(doc)


def _add_heading2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_paragraph_font(p, text, 13, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _add_heading3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_paragraph_font(p, text, 11, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _add_body(doc: Document, text: str, indent: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    run.font.name = FONT_NAME_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    run.font.size = Pt(10.5)


def _fmt(value: Any, digits: int = 4, fallback: str = "—") -> str:
    """将数值格式化为字符串，保留指定小数位。"""
    if value is None:
        return fallback
    try:
        f = float(value)
        if not (-1e30 < f < 1e30):
            return fallback
        return f"{f:.{digits}f}"
    except (TypeError, ValueError):
        return fallback


def _fmt_percent(value: Any, digits: int = 2) -> str:
    if value is None:
        return "—"
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except (TypeError, ValueError):
        return "—" if value is None else str(value)


def _fmt_ci95(ci: dict[str, Any] | None, prefix: str = "") -> str:
    """格式化 95% 置信区间为字符串。"""
    if not ci or not ci.get("available"):
        return "—"
    lower = ci.get("lower")
    upper = ci.get("upper")
    if lower is None or upper is None:
        return "—"
    return f"[{_fmt(lower, 4)} ~ {_fmt(upper, 4)}]"


# ---------------------------------------------------------------------------
# 各章节生成函数
# ---------------------------------------------------------------------------

def _build_cover(doc: Document, task: RunTask) -> None:
    """第一章：封面。"""
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("SEM 图像分割统计实验报告")
    run.bold = True
    run.font.name = FONT_NAME_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("面向扫描电镜图像的镍基单晶微观结构分割统计系统")
    sub_run.font.name = FONT_NAME_EN
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    doc.add_paragraph()

    # 信息表
    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = "Table Grid"
    info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    created_at = ""
    raw_created = getattr(task, 'created_at', None) or getattr(task, 'started_at', None)
    if raw_created:
        try:
            dt = raw_created
            created_at = dt.strftime("%Y-%m-%d %H:%M") if hasattr(dt, "strftime") else str(dt)
        except Exception:
            created_at = str(task.created_at)

    meta = [
        ("任务名称", task.name or "—"),
        ("创建时间", created_at),
        ("分割模式", "传统分割" if task.segmentation_mode == "traditional" else "深度学习"),
        ("输入模式", "批量" if task.input_mode == "batch" else "单张"),
    ]
    for col_idx, (key, val) in enumerate(meta):
        cell = info_table.rows[0].cells[col_idx]
        cell.width = Cm(3.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = cell.paragraphs[0]
        r1 = p.add_run(f"{key}\n")
        r1.bold = True
        r1.font.name = FONT_NAME_EN
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
        r1.font.size = Pt(9)
        r2 = p.add_run(val)
        r2.font.name = FONT_NAME_EN
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
        r2.font.size = Pt(9)
        _set_cell_bg(cell, "DEEAF1")

    doc.add_page_break()


def _build_overview(doc: Document, task: RunTask, config: dict[str, Any]) -> None:
    """第二章：任务概述。"""
    _add_heading1(doc, "一、任务概述")

    summary = task.summary or {}

    # 基本信息
    _add_heading2(doc, "1.1 基本信息")
    rows = [
        ("任务名称", task.name or "—"),
        ("分割模式", "传统分割" if task.segmentation_mode == "traditional" else "深度学习"),
        ("输入模式", "批量" if task.input_mode == "batch" else "单张"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
        for cell in tbl.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT_NAME_EN
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
                    r.font.size = Pt(10)
        _set_cell_bg(tbl.rows[i].cells[0], "F2F2F2")
        for r in tbl.rows[i].cells[0].paragraphs[0].runs:
            r.bold = True
    tbl.columns[0].width = Cm(4)
    tbl.columns[1].width = Cm(12)

    # 图像数量
    batch = summary.get("batch") or {}
    total_images = 0
    for mode_data in batch.values():
        if isinstance(mode_data, dict):
            total_images = max(total_images, int(mode_data.get("image_count", 0)))

    # 预处理配置
    _add_heading2(doc, "1.2 预处理配置")
    preprocess = config.get("preprocess") or {}
    pp_enabled = preprocess.get("enabled", False)
    _add_body(doc, f"预处理状态：{'已启用' if pp_enabled else '未启用'}")
    if pp_enabled:
        ops = []
        bg = preprocess.get("background") or {}
        if bg.get("method", "none") != "none":
            ops.append(f"背景校正（{bg.get('method')}，半径={bg.get('radius', 25)}）")
        dn = preprocess.get("denoise") or {}
        if dn.get("method", "none") != "none":
            ops.append(f"去噪（{dn.get('method')}）")
        eh = preprocess.get("enhance") or {}
        if eh.get("method", "none") != "none":
            ops.append(f"对比度增强（{eh.get('method')}）")
        if ops:
            for op in ops:
                _add_body(doc, f"  • {op}", indent=True)
        else:
            _add_body(doc, "  （预处理已启用但未配置具体操作）")

    # 分割配置
    _add_heading2(doc, "1.3 分割配置")
    if task.segmentation_mode == "traditional":
        seg = config.get("traditional_seg") or {}
        method_map = {
            "threshold": "阈值分割",
            "adaptive": "自适应阈值分割",
            "edge": "边缘检测分割",
            "clustering": "聚类分割",
        }
        seg_method = method_map.get(seg.get("method", ""), seg.get("method", "—"))
        _add_body(doc, f"分割方法：{seg_method}")
        _add_body(doc, f"目标相：{'暗相（γ 基体）' if seg.get('foreground_target') == 'dark' else '亮相（γ′ 强化相）'}")
        if seg.get("method") == "threshold":
            mode_t = seg.get("threshold_mode", "otsu")
            mode_label = {"otsu": "Otsu 自动阈值", "global": "全局固定阈值", "fixed": "手动固定阈值"}.get(mode_t, mode_t)
            _add_body(doc, f"阈值模式：{mode_label}")
            if mode_t == "global":
                _add_body(doc, f"全局阈值：{seg.get('global_threshold', 120)}")
        if seg.get("method") == "adaptive":
            _add_body(doc, f"邻域计算方式：{'高斯' if seg.get('adaptive_method') == 'gaussian' else '均值'}")
            _add_body(doc, f"块大小：{seg.get('adaptive_block_size', 35)}，偏移量 C：{seg.get('adaptive_c', 5)}")
        _add_body(doc, f"最小面积过滤：{seg.get('min_area', 30)} px")
        if seg.get("fill_holes"):
            _add_body(doc, "填补空洞：已启用")
        if seg.get("watershed"):
            _add_body(doc, f"Watershed 分离：已启用（分离度={seg.get('watershed_separation', 35)}）")
        if seg.get("remove_border"):
            _add_body(doc, "触边剔除：已启用")
    else:
        dl = config.get("dl_model") or {}
        slot = dl.get("model_slot", "—")
        slot_label = {
            "mbu_netpp": "MBU-Net++",
            "sam_lora": "SAM LoRA",
            "resnext50": "ResNeXt50",
            "matsam": "MatSAM",
            "custom": "自定义模型",
        }.get(slot, slot)
        device = dl.get("device", "auto")
        _add_body(doc, f"模型槽位：{slot_label}")
        _add_body(doc, f"运行设备：{device}")
        if dl.get("input_size"):
            _add_body(doc, f"输入尺寸：{dl.get('input_size')} px（滑窗）")

    # 标定信息
    _add_heading2(doc, "1.4 标定信息")
    calibration_probe = summary.get("calibration_probe")
    if calibration_probe and isinstance(calibration_probe, dict):
        ocr_info = []
        if calibration_probe.get("ocr_scale_bar_um"):
            ocr_info.append(f"比例尺：{calibration_probe['ocr_scale_bar_um']} μm")
        if calibration_probe.get("ocr_magnification_text"):
            ocr_info.append(f"放大倍率：{calibration_probe['ocr_magnification_text']}")
        if calibration_probe.get("ocr_fov_um"):
            ocr_info.append(f"视场：{calibration_probe['ocr_fov_um']} μm")
        if calibration_probe.get("ocr_wd_mm"):
            ocr_info.append(f"工作距离：{calibration_probe['ocr_wd_mm']} mm")
        if calibration_probe.get("ocr_detector"):
            ocr_info.append(f"探测器：{calibration_probe['ocr_detector']}")
        if calibration_probe.get("ocr_scan_mode"):
            ocr_info.append(f"扫描模式：{calibration_probe['ocr_scan_mode']}")
        if ocr_info:
            for info in ocr_info:
                _add_body(doc, f"  • {info}", indent=True)
        else:
            _add_body(doc, "  （OCR 未成功解析，具体参数请参见配置快照）")
    else:
        cal_hint = summary.get("calibration_hint")
        if cal_hint:
            _add_body(doc, f"  {cal_hint}")
        else:
            _add_body(doc, "  未检测到标定信息，单位以像素计")

    doc.add_page_break()


def _build_image_results(
    doc: Document,
    task: RunTask,
    metrics: list[MetricRecord],
) -> None:
    """
    逐图分割结果：原图 + 分割叠加对比图。
    放在统计结果汇总之前，让读者先有直观认识。
    """
    _add_heading1(doc, "二、逐图分割结果")

    if not task.images:
        _add_body(doc, "当前任务无输入图像。")
        return

    metrics_by_image = {m.image_id: m for m in metrics if m.image_id is not None}

    for idx, img in enumerate(task.images, 1):
        mode = task.segmentation_mode
        mode_metrics = [
            m for m in metrics
            if m.image_id == img.id and m.mode == mode
        ]
        if not mode_metrics:
            mode_metrics = [m for m in metrics if m.image_id == img.id]
        metric = mode_metrics[0] if mode_metrics else None

        _add_heading2(doc, f"{idx}. {img.original_name}")

        img_path = storage_service.absolute_path(img.relative_path)
        overlay_path = None
        mask_path = None
        object_overlay_path = None

        if metric and metric.artifacts:
            art = metric.artifacts
            overlay_path = storage_service.absolute_path(art.get("overlay_path", ""))
            mask_path = storage_service.absolute_path(art.get("mask_path", ""))
            object_overlay_path = storage_service.absolute_path(art.get("object_overlay_path", ""))

        available_paths = {
            "原图": img_path,
            "分割叠加": overlay_path,
            "掩码": mask_path,
            "对象叠加": object_overlay_path,
        }

        present = {k: v for k, v in available_paths.items() if v and v.exists()}
        if not present:
            _add_body(doc, f"  （图像分割结果文件不存在，请确认任务已完整执行）")
            continue

        # 根据有几张图决定布局
        keys = list(present.keys())
        cols = len(keys)
        cols = min(cols, 4)

        # 创建一个 1 行 N 列的表格，每格放一张图
        tbl = doc.add_table(rows=1, cols=cols)
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for ci, (label, path) in enumerate(list(present.items())[:cols]):
            cell = tbl.rows[0].cells[ci]
            cell.width = Cm(6.5 if cols <= 2 else (4.0 if cols <= 4 else 3.0))

            try:
                cell.paragraphs[0].add_run().add_picture(str(path), width=Cm(6.0 if cols <= 2 else (3.8 if cols <= 4 else 2.8)))
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(label)
                r.font.name = FONT_NAME_EN
                r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
                r.font.size = Pt(8)
                r.bold = True
            except Exception:
                cell.paragraphs[0].add_run(f"[{label}]").font.size = Pt(8)

        doc.add_paragraph()


def _fmt_vf_row(
    summary: dict[str, Any],
    image_names: list[str],
    area_unit: str,
    diameter_unit: str,
) -> list[list[str]]:
    """生成一张体积分数表格的数据行。"""
    vf_values = summary.get("volume_fractions", [])
    area_values = summary.get("image_mean_areas", [])
    size_values = summary.get("image_mean_sizes", [])
    count_values = summary.get("particle_counts", summary.get("object_counts", []))
    ci = summary.get("volume_fraction_ci95")

    rows = []
    for i, img_name in enumerate(image_names):
        vf = vf_values[i] if i < len(vf_values) else None
        area = area_values[i] if i < len(area_values) else None
        size = size_values[i] if i < len(size_values) else None
        count = count_values[i] if i < len(count_values) else None
        rows.append([
            img_name,
            _fmt_percent(vf, 2),
            _fmt(count, 0, "—"),
            _fmt(area, 4),
            area_unit,
            _fmt(size, 4),
            diameter_unit,
        ])
    return rows


def _build_statistics(doc: Document, task: RunTask) -> None:
    """第三章：统计结果汇总。"""
    _add_heading1(doc, "三、统计结果汇总")

    summary = task.summary or {}
    batch = summary.get("batch") or {}

    if not batch:
        _add_body(doc, "当前任务未生成统计结果。")
        return

    mode_labels = {
        "traditional": "传统分割",
        "dl": "深度学习分割",
    }

    for mode, mode_data in batch.items():
        if not isinstance(mode_data, dict):
            continue

        mode_label = mode_labels.get(mode, mode)
        _add_heading2(doc, f"2.{'1' if mode == 'traditional' else '2'} {mode_label}")

        image_names = mode_data.get("image_names", [])
        area_unit = mode_data.get("area_unit", "px²")
        size_unit = mode_data.get("size_unit", "px")
        diameter_unit = mode_data.get("diameter_unit", "px")
        channel_width_unit = mode_data.get("channel_width_unit", "px")

        image_count = mode_data.get("image_count", 0)

        # 2.x.1 体积分数
        _add_heading3(doc, f"2.{'1' if mode == 'traditional' else '2'}.1 体积分数（VF）")
        _add_body(doc, f"共 {image_count} 张图像参与统计。")

        if image_names:
            headers = ["图像名称", "VF (%)", "颗粒计数", f"平均面积 ({area_unit})", "面积单位", f"平均等效直径 ({diameter_unit})", "直径单位"]
            col_widths = [Cm(3.5), Cm(1.8), Cm(1.8), Cm(2.5), Cm(1.5), Cm(2.5), Cm(1.5)]
            tbl = doc.add_table(rows=1 + len(image_names), cols=len(headers))
            tbl.style = "Table Grid"

            # 表头
            hrow = tbl.rows[0]
            for ci, (h, w) in enumerate(zip(headers, col_widths)):
                hrow.cells[ci].text = h
                hrow.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in hrow.cells[ci].paragraphs[0].runs:
                    r.bold = True
                    r.font.name = FONT_NAME_EN
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(hrow.cells[ci], "2E75B6")
                hrow.cells[ci].width = w

            # 数据行
            vf_values = mode_data.get("volume_fractions", [])
            area_values = mode_data.get("image_mean_areas", [])
            size_values = mode_data.get("image_mean_sizes", [])
            count_values = mode_data.get("particle_counts", mode_data.get("object_counts", []))

            for ri, img_name in enumerate(image_names):
                drow = tbl.rows[ri + 1]
                vf = vf_values[ri] if ri < len(vf_values) else None
                area = area_values[ri] if ri < len(area_values) else None
                size = size_values[ri] if ri < len(size_values) else None
                count = count_values[ri] if ri < len(count_values) else None

                cells_data = [
                    img_name,
                    _fmt_percent(vf, 2),
                    _fmt(count, 0, "—"),
                    _fmt(area, 4),
                    area_unit,
                    _fmt(size, 4),
                    diameter_unit,
                ]
                for ci, val in enumerate(cells_data):
                    drow.cells[ci].text = val
                    drow.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in drow.cells[ci].paragraphs[0].runs:
                        r.font.name = FONT_NAME_EN
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
                        r.font.size = Pt(9)
                    if ri % 2 == 1:
                        _set_cell_bg(drow.cells[ci], "F2F2F2")
        else:
            _add_body(doc, "（无有效统计数据）")

        # 2.x.2 批量汇总
        _add_heading3(doc, f"2.{'1' if mode == 'traditional' else '2'}.2 批量统计汇总")
        ci_summary = mode_data.get("volume_fraction_ci95")
        if ci_summary and ci_summary.get("available"):
            ci_lower = ci_summary.get("lower")
            ci_upper = ci_summary.get("upper")
            ci_text = f"VF = {_fmt_percent(mode_data.get('volume_fraction', 0), 2)}，95% CI: [{_fmt_percent(ci_lower, 2)} ~ {_fmt_percent(ci_upper, 2)}]，n = {ci_summary.get('n', '—')}"
            _add_body(doc, ci_text)
        else:
            _add_body(doc, f"VF = {_fmt_percent(mode_data.get('volume_fraction', 0), 2)}（样本量不足，无法计算置信区间）")

        _add_body(doc, f"颗粒总数：{_fmt(mode_data.get('particle_count_total', 0), 0)}")

        # 尺寸分布
        if mode_data.get("areas"):
            _add_heading3(doc, f"2.{'1' if mode == 'traditional' else '2'}.3 尺寸分布统计（面积）")
            size_rows = [
                ["统计量", "面积均值", "面积中位数", "面积标准差", "最小值", "最大值", "单位"],
                ["数值",
                    _fmt(mode_data.get("mean_area"), 4),
                    _fmt(mode_data.get("median_area"), 4),
                    _fmt(mode_data.get("std_area"), 4),
                    _fmt(min(mode_data["areas"]), 4) if mode_data["areas"] else "—",
                    _fmt(max(mode_data["areas"]), 4) if mode_data["areas"] else "—",
                    area_unit,
                ],
            ]
            tbl2 = doc.add_table(rows=len(size_rows), cols=len(size_rows[0]))
            tbl2.style = "Table Grid"
            for ri, row_data in enumerate(size_rows):
                for ci, val in enumerate(row_data):
                    tbl2.rows[ri].cells[ci].text = val
                    tbl2.rows[ri].cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in tbl2.rows[ri].cells[ci].paragraphs[0].runs:
                        r.font.name = FONT_NAME_EN
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
                        r.font.size = Pt(9)
                        r.bold = (ri == 0)
                    if ri == 0:
                        for c in tbl2.rows[ri].cells:
                            _set_cell_bg(c, "2E75B6")
                            for r in c.paragraphs[0].runs:
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # 通道宽度（如果有）
        cw_x = mode_data.get("mean_channel_width_x")
        cw_y = mode_data.get("mean_channel_width_y")
        if cw_x is not None or cw_y is not None:
            _add_heading3(doc, f"2.{'1' if mode == 'traditional' else '2'}.4 通道宽度测量")
            _add_body(doc, "测量方法：距离变换骨架法（旋转不变）")
            if cw_x is not None:
                _add_body(doc, f"  X 方向通道宽度均值：{_fmt(cw_x, 4)} {channel_width_unit}")
            if cw_y is not None:
                _add_body(doc, f"  Y 方向通道宽度均值：{_fmt(cw_y, 4)} {channel_width_unit}")

        doc.add_page_break()


def _build_charts(doc: Document, task: RunTask, chart_paths: dict[str, str]) -> None:
    """第四章：图表。"""
    _add_heading1(doc, "四、统计图表")

    summary = task.summary or {}
    batch = summary.get("batch") or {}
    chart_mode_order = ["traditional", "dl"]

    if not chart_paths:
        _add_body(doc, "当前任务未生成图表。")

    for mode in chart_mode_order:
        mode_label = {"traditional": "传统分割", "dl": "深度学习分割"}.get(mode, mode)
        if mode not in batch or not isinstance(batch[mode], dict):
            continue

        mode_data = batch[mode]
        area_key = f"{mode}_area_hist"
        size_key = f"{mode}_size_hist"
        vf_key = f"{mode}_vf_bar"

        area_rel = chart_paths.get(area_key)
        size_rel = chart_paths.get(size_key)
        vf_rel = chart_paths.get(vf_key)

        if not area_rel and not size_rel and not vf_rel:
            continue

        _add_heading2(doc, f"5.{'1' if mode == 'traditional' else '2'} {mode_label}")

        if area_rel:
            _add_heading3(doc, "面积分布直方图")
            abs_path = storage_service.absolute_path(area_rel)
            if abs_path.exists():
                try:
                    doc.add_picture(str(abs_path), width=Inches(6.0))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    _add_body(doc, "  （图表文件读取失败）")
            else:
                _add_body(doc, "  （图表文件不存在）")

        if size_rel:
            _add_heading3(doc, "等效直径分布直方图")
            abs_path = storage_service.absolute_path(size_rel)
            if abs_path.exists():
                try:
                    doc.add_picture(str(abs_path), width=Inches(6.0))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    _add_body(doc, "  （图表文件读取失败）")
            else:
                _add_body(doc, "  （图表文件不存在）")

        if vf_rel:
            _add_heading3(doc, "体积分数柱状图")
            abs_path = storage_service.absolute_path(vf_rel)
            if abs_path.exists():
                try:
                    doc.add_picture(str(abs_path), width=Inches(6.0))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    _add_body(doc, "  （图表文件读取失败）")
            else:
                _add_body(doc, "  （图表文件不存在）")

        doc.add_page_break()


def _build_particles(doc: Document, task: RunTask) -> None:
    """第五章：逐颗粒明细。"""
    _add_heading1(doc, "五、逐颗粒明细")

    summary = task.summary or {}
    batch = summary.get("batch") or {}

    if not batch:
        _add_body(doc, "当前任务无颗粒明细数据。")
        return

    for mode, mode_data in batch.items():
        if not isinstance(mode_data, dict):
            continue

        mode_label = {"traditional": "传统分割", "dl": "深度学习分割"}.get(mode, mode)
        _add_heading2(doc, f"5.{'1' if mode == 'traditional' else '2'} {mode_label}")

        areas = mode_data.get("areas", [])
        diameters = mode_data.get("diameters", [])

        if not areas:
            _add_body(doc, "（无有效颗粒数据）")
            continue

        area_unit = mode_data.get("area_unit", "px²")
        diameter_unit = mode_data.get("diameter_unit", "px")

        # 截断：Word 表格行数过多会导致文件过大，只保留前 500 行
        MAX_ROWS = 500
        display_count = min(len(areas), MAX_ROWS)

        headers = [
            "序号", f"面积 ({area_unit})", f"等效直径 ({diameter_unit})",
            "圆度", "球度", "实心度", "长轴", "短轴", "过滤"
        ]
        col_widths = [Cm(1.2), Cm(2.5), Cm(2.5), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8)]

        tbl = doc.add_table(rows=1 + display_count, cols=len(headers))
        tbl.style = "Table Grid"

        # 表头
        hrow = tbl.rows[0]
        for ci, (h, w) in enumerate(zip(headers, col_widths)):
            hrow.cells[ci].text = h
            hrow.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hrow.cells[ci].width = w
            for r in hrow.cells[ci].paragraphs[0].runs:
                r.bold = True
                r.font.name = FONT_NAME_EN
                r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(hrow.cells[ci], "2E75B6")

        for ri in range(display_count):
            drow = tbl.rows[ri + 1]
            area = areas[ri] if ri < len(areas) else None
            diam = diameters[ri] if ri < len(diameters) else None
            row_data = [
                str(ri + 1),
                _fmt(area, 4),
                _fmt(diam, 4),
                _fmt(mode_data.get("circularities", [None] * ri)[ri] if ri < len(mode_data.get("circularities", [])) else None, 4),
                _fmt(mode_data.get("solidities", [None] * ri)[ri] if ri < len(mode_data.get("solidities", [])) else None, 4),
                _fmt(mode_data.get("roundnesses", [None] * ri)[ri] if ri < len(mode_data.get("roundnesses", [])) else None, 4),
                _fmt(mode_data.get("majors", [None] * ri)[ri] if ri < len(mode_data.get("majors", [])) else None, 4),
                _fmt(mode_data.get("minors", [None] * ri)[ri] if ri < len(mode_data.get("minors", [])) else None, 4),
                "—",
            ]
            for ci, val in enumerate(row_data):
                drow.cells[ci].text = val
                drow.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in drow.cells[ci].paragraphs[0].runs:
                    r.font.name = FONT_NAME_EN
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CN)
                    r.font.size = Pt(8)
                if ri % 2 == 1:
                    _set_cell_bg(drow.cells[ci], "F2F2F2")

        if len(areas) > MAX_ROWS:
            doc.add_paragraph()
            _add_body(doc, f"（共 {len(areas)} 个颗粒，仅展示前 {MAX_ROWS} 条记录，完整数据见 particles.xlsx）")

        doc.add_page_break()


def _build_config_snapshot(doc: Document, config: dict[str, Any]) -> None:
    """第六章：实验配置快照。"""
    _add_heading1(doc, "五、实验配置快照")

    _add_body(doc, "以下为本次任务的完整配置参数，以 JSON 格式保存。")

    json_text = json.dumps(config, ensure_ascii=False, indent=2, default=str)
    # Word 中插入 JSON 块（等宽字体）
    p = doc.add_paragraph()
    run = p.add_run(json_text[:8000])  # 截断避免过长
    run.font.name = "Courier New"
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------

def generate_docx_report(
    db: Session,
    run: RunTask,
    metrics: list[MetricRecord],
    config_snapshot: dict[str, Any],
    chart_paths: dict[str, str],
) -> str:
    """
    生成 Word 格式实验报告。

    Returns:
        报告文件的相对路径字符串（相对于 storage/ 目录）。
    """
    doc = Document()

    # 页面设置：横向 A4，页边距
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)

    # 封面
    _build_cover(doc, run)

    # 任务概述
    _build_overview(doc, run, config_snapshot)

    # 逐图分割结果
    _build_image_results(doc, run, metrics)

    # 统计结果
    _build_statistics(doc, run)

    # 图表
    _build_charts(doc, run, chart_paths)

    # 逐颗粒明细
    _build_particles(doc, run)

    # 配置快照
    _build_config_snapshot(doc, config_snapshot)

    # 保存文件
    export_dir = storage_service.run_subdir(run.id, "exports")
    filename = f"experiment_report_{run.id}.docx"
    output_path = export_dir / filename
    doc.save(str(output_path))

    return storage_service.relative_path(output_path)
